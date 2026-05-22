"""Tests for file parsers (PDF, DOCX, DOC)."""

import pytest
from pathlib import Path


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Create a minimal PDF with text content."""
    import fitz
    path = tmp_path / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 72), "Hello PDF World")
    page.insert_text((50, 100), "Paragraph Two")
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX with paragraph and table."""
    from docx import Document
    path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello DOCX World")
    doc.add_paragraph("这是第二段中文内容。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    doc.save(str(path))
    return path


@pytest.fixture
def empty_pdf(tmp_path: Path) -> Path:
    """Create a PDF with no text (blank pages)."""
    import fitz
    path = tmp_path / "empty.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def legacy_doc(tmp_path: Path) -> Path:
    """Create a fake .doc file that is not OOXML (binary garbage)."""
    path = tmp_path / "legacy.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 100)
    return path


@pytest.mark.asyncio
async def test_extract_text_from_pdf(sample_pdf):
    """PDF parsing should return concatenated page text."""
    from gustobot.infrastructure.file_parsers import extract_text_from_pdf
    text = await extract_text_from_pdf(sample_pdf)
    assert "Hello PDF World" in text
    assert "Paragraph Two" in text


@pytest.mark.asyncio
async def test_extract_text_from_docx(sample_docx):
    """DOCX parsing should return paragraph and table text."""
    from gustobot.infrastructure.file_parsers import extract_text_from_docx
    text = await extract_text_from_docx(sample_docx)
    assert "Hello DOCX World" in text
    assert "中文内容" in text
    assert "A1 | B1" in text or "A1" in text


@pytest.mark.asyncio
async def test_extract_text_from_pdf_empty(empty_pdf):
    """Empty PDF (no text) should raise ValueError."""
    from gustobot.infrastructure.file_parsers import extract_text_from_pdf
    with pytest.raises(ValueError, match="未能提取到文本内容"):
        await extract_text_from_pdf(empty_pdf)


@pytest.mark.asyncio
async def test_extract_text_from_doc_legacy(legacy_doc):
    """Legacy .doc binary file should raise ValueError with conversion hint."""
    from gustobot.infrastructure.file_parsers import extract_text_from_doc
    with pytest.raises(ValueError, match="暂不支持旧版"):
        await extract_text_from_doc(legacy_doc)


@pytest.mark.asyncio
async def test_extract_text_from_nonexistent_file():
    """Non-existent file should propagate an error."""
    import fitz
    from gustobot.infrastructure.file_parsers import extract_text_from_pdf
    with pytest.raises((FileNotFoundError, fitz.FileNotFoundError)):
        await extract_text_from_pdf(Path("/nonexistent/file.pdf"))
