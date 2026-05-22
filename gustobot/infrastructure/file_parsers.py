"""Text extraction from uploaded file formats (PDF, DOCX, DOC, etc.)"""

import asyncio
from pathlib import Path

from loguru import logger


async def extract_text_from_pdf(path: Path) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        import fitz
    except ImportError:
        raise RuntimeError("PyMuPDF (fitz) is not installed. Cannot parse PDF files.")

    def _extract() -> str:
        doc = fitz.open(str(path))
        try:
            pages = []
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                pages.append(text)
            return "\n".join(pages)
        finally:
            doc.close()

    text = await asyncio.to_thread(_extract)
    if not text.strip():
        raise ValueError("PDF文件未能提取到文本内容（可能是扫描件或图片型PDF）。")
    return text


async def extract_text_from_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx.

    Extracts paragraph text plus text from tables.
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is not installed. Cannot parse DOCX files.")

    def _extract() -> str:
        doc = Document(str(path))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))

        return "\n".join(parts)

    text = await asyncio.to_thread(_extract)
    if not text.strip():
        raise ValueError("DOCX文件未能提取到文本内容（文件可能为空）。")
    return text


async def extract_text_from_doc(path: Path) -> str:
    """Extract text from a .doc file.

    python-docx only handles .docx (OOXML) format. Legacy .doc (OLE binary)
    files are not supported. We attempt to open as docx first since many
    modern .doc files are actually OOXML with a .doc extension.
    """
    try:
        return await extract_text_from_docx(path)
    except Exception:
        raise ValueError(
            "暂不支持旧版 .doc 格式（OLE二进制格式）。请将文件另存为 .docx 或 .pdf 后重新上传。"
        )
